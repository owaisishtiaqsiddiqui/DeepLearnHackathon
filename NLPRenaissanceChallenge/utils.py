# utils.py — RenAIssance OCR Utility Functions
import fitz
import cv2
import os
from docx import Document
import numpy as np
from PIL import Image, ImageOps
from pathlib import Path
import pandas as pd
from natsort import natsorted
import string
import re
from natsort import natsorted

# ------------------ File and Text Utilities ------------------
def count_files_in_folder(folder_path, extensions_list):
    # Initialize counter for files
    file_count = 0

    # Iterate through all files in the folder
    for filename in os.listdir(folder_path):
        # Check if the file ends with the given file extension
        for extension in extensions_list:
            if filename.lower().endswith(extension):
                file_count += 1

    return file_count
def remove_punctuation(text):
    # Remove anything like (1), (23), etc.
    text = re.sub(r'\(\d+\)', '', text)
    # Remove punctuation
    return text.translate(str.maketrans('', '', string.punctuation)).strip()
def count_occurrences_of_semicolon(filename):
    with open(filename, 'r', encoding="utf-8") as file:
        content = file.read()
        return content.count(';')
def read_nth_line(file_path, n):
    """Read the nth (0-based) line from a file (UTF-8 safe)."""
    with open(file_path, 'r', encoding="utf-8") as file:
        for i, line in enumerate(file):
            if i == n:
                return line.strip()
    return None
def save_pages_to_text(docx_file, output_file):
    document = Document(str(docx_file))
    all_text = ""
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        # Remove punctuations from the text before appending to the all_text variable
        text_without_punctuation = remove_punctuation(text)
        # Check if the line starts with "PDF p"
        if not text.startswith("PDF p"):
            all_text += text + "\n"

    # Write all_text to the output file
    with open(output_file, "w") as file:
        file.write(all_text)

# ------------------ PDF and Image Processing ------------------
def pad_and_resize_images(folder_path):
    """
    Pads and resizes images in a folder (recursively) to a fixed 4:1 aspect ratio and target size.
    Skips malformed images and logs a summary.
    """
    if not os.path.exists(folder_path):
        raise ValueError(f"The folder {folder_path} does not exist")

    target_aspect_ratio = 4  # Width:Height
    target_width = 200
    target_height = 40

    processed_count = 0
    skipped_files = []
    total_files = 0

    for root, _, files in os.walk(folder_path):
        for filename in files:
            file_path = os.path.join(root, filename)

            if not filename.lower().endswith((".png", ".jpg", ".jpeg")):
                continue

            total_files += 1
            try:
                with Image.open(file_path) as img:
                    img = img.convert('L')  # Convert to grayscale
                    width, height = img.size
                    aspect_ratio = width / height

                    # Pad to target aspect ratio
                    if aspect_ratio < target_aspect_ratio:
                        new_width = int(height * target_aspect_ratio)
                        padding = max((new_width - width) // 2, 0)
                        padded_img = ImageOps.expand(img, border=(padding, 0, padding, 0), fill='white')
                    else:
                        padded_img = img

                    # Resize
                    resized_img = padded_img.resize((target_width, target_height), Image.ANTIALIAS)
                    resized_img.save(file_path)
                    processed_count += 1
                    print(f"Processed: {file_path}")

            except Exception as e:
                print(f"Skipped: {file_path} — {e}")
                skipped_files.append(file_path)

    # --- Summary Report ---
    print("\n--- Summary ---")
    print(f"Total images found       : {total_files}")
    print(f"Successfully processed   : {processed_count}")
    print(f"Skipped due to errors    : {len(skipped_files)}")

    if skipped_files:
        print("\nSkipped files:")
        for f in skipped_files:
            print(f"- {f}")

def pdf_to_images(pdf_path, output_folder):
    pdf_document = fitz.open(pdf_path)
def split_and_save_image(image_path, output_folder, last_image_number):
    # Read the image
    img = cv2.imread(image_path)

    # Get image width
    _, width, _ = img.shape

    # You can alter these, to get the optimum values for both
    width_for_single_page, width_for_dual_pages = 350, 450

    # Determine filename based on width and last image number
    if width < width_for_single_page:
        filename = f"image_{last_image_number}.png"
        output_path = os.path.join(output_folder, filename)
        cv2.imwrite(output_path, img)
        last_image_number += 1

    elif width > width_for_dual_pages:
        left_half = img[:, :width // 2]
        right_half = img[:, width // 2:]
        filename = f"image_{last_image_number}.png"
        output_path = os.path.join(output_folder, filename)
        cv2.imwrite(output_path, left_half)
        last_image_number += 1
        filename = f"image_{last_image_number}.png"
        output_path = os.path.join(output_folder, filename)
        cv2.imwrite(output_path, right_half)
        last_image_number += 1

    return last_image_number
def process_images(image_folder, output_folder):
    # Initialize image counter, starting from 1
    last_image_number = 1

    # Iterate through all files in the image folder
    for indx in range(count_files_in_folder(image_folder, [".png", ".jpg", ".jpeg"])):
        # Check for image files only
        filename = 'page_' + str(indx+1) + '.png'
        image_path = os.path.join(image_folder, filename)
        last_image_number = split_and_save_image(image_path, output_folder, last_image_number)
def resize_images_recursively(input_root, output_root, new_size=(200, 50)):
    input_root = Path(input_root)
    output_root = Path(output_root)
    output_root.mkdir(parents=True, exist_ok=True)

    # Traverse subdirectories (datasets)
    for subdir in input_root.iterdir():
        if subdir.is_dir():
            output_subdir = output_root / subdir.name
            output_subdir.mkdir(parents=True, exist_ok=True)

            for filename in os.listdir(subdir):
                input_path = subdir / filename
                if input_path.suffix.lower() in ['.png', '.jpg', '.jpeg']:
                    try:
                        with Image.open(input_path) as img:
                            resized = img.resize(new_size)
                            output_path = output_subdir / (input_path.stem + ".png")
                            resized.save(output_path)
                    except Exception as e:
                        print(f"Could not process {input_path}: {e}")

    print(f"All images resized to {new_size} and saved in: {output_root}")

# ------------------ Image Enhancement ------------------
def load_image(self, img_path):
        img = cv2.imread(img_path, cv2.IMREAD_COLOR)
        if img is None:
            print(f"Could not read image: {img_path}")
        return img
def denoise(img):
    return cv2.fastNlMeansDenoising(img, None, h=5, templateWindowSize=7, searchWindowSize=21)
def remove_background(img):
    bg = cv2.medianBlur(img, 11)
    return cv2.divide(img, bg, scale=255)
def enhance_text(img):
    clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
    enhanced = clahe.apply(img)
    lo, hi = np.percentile(enhanced, [2, 98])
    stretched = np.clip((enhanced - lo) * (255.0 / (hi - lo)), 0, 255).astype(np.uint8)
    return stretched
def brighten_whites(img, gamma=0.6):
    inv_gamma = 1.0 / gamma
    table = np.array([(i / 255.0) ** inv_gamma * 255 for i in range(256)]).astype("uint8")
    return cv2.LUT(img, table)
def add_blue_tint(gray_img, intensity=50):
    bgr = cv2.cvtColor(gray_img, cv2.COLOR_GRAY2BGR)
    b, g, r = cv2.split(bgr)
    b = cv2.add(b, np.full_like(b, intensity))
    return cv2.merge((b, g, r))
def preprocess(img_bgr):
    gray = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY)
    gray_dn = denoise(gray)
    bg_removed = remove_background(gray_dn)
    bw_enhanced = enhance_text(bg_removed)
    inverted = cv2.bitwise_not(bw_enhanced)
    white = brighten_whites(inverted)
    blue = add_blue_tint(white)
    return blue
def run_full_preprocessing_pipeline(
    pdf_path,
    output_raw_folder,
    output_split_folder,
    output_preprocessed_folder,
    source_suffix):

    # Ensure all output folders exist
    Path(output_raw_folder).mkdir(parents=True, exist_ok=True)
    Path(output_split_folder).mkdir(parents=True, exist_ok=True)
    Path(output_preprocessed_folder).mkdir(parents=True, exist_ok=True)

    # 1. Convert PDF to raw page images
    print("Converting PDF to images...")
    pdf_to_images(pdf_path, output_raw_folder)

    # 2. Split double-page scans
    print("Splitting double-page images...")
    process_images(output_raw_folder, output_split_folder)

    # 3. Preprocess each split image
    print("Preprocessing images...")
    for img_path in sorted(Path(output_split_folder).glob("*.png")):
        img = cv2.imread(str(img_path))
        if img is None:
            print(f"Could not read image: {img_path}")
            continue
        stem = img_path.stem
        new_name = f"{stem}_{source_suffix}.png"
        output_path = Path(output_preprocessed_folder) / new_name
        processed = preprocess(img)
        cv2.imwrite(str(output_path), processed)
        print(f"Saved preprocessed image: {output_path.name}")

    print("Full preprocessing pipeline complete.")

# ------------------ Ground Truth & Transcription ------------------
def extract_ground_truth_for_dataset(
    docx_file, 
    output_path="groundTruth.txt"
):
    """
    Extracts and saves ground truth text from a single .docx file to a text file.
    """
    docx_file = Path(docx_file)
    if not docx_file.exists():
        print(f".docx not found: {docx_file}")
        return

    document = Document(str(docx_file))
    extracted_text = ""
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text: # and not text.startswith("PDF p"):
            extracted_text += text + "\n"

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as file:
        file.write(extracted_text)
    print(f"Ground truth text saved to {output_path}")

# ------------------ Bounding Box Utilities ------------------
def process_bounding_boxes(file_path):
    with open(file_path, "r", encoding="utf-8") as file:
        lines = file.readlines()

    bounding_boxes = []
    for line in lines:
        line = line.strip()
        if not line or ";" in line:
            continue  # Skip empty lines or semicolon-only lines
        parts = line.split(',')
        if len(parts) != 8:
            print(f"Skipping malformed box: {line}")
            continue
        try:
            coords = list(map(int, parts))
            bounding_boxes.append(coords)
        except ValueError:
            print(f"Non-integer values in box: {line}")
            continue

    # Sort by top y-coord
    bounding_boxes.sort(key=lambda box: box[1])

    vertical_distance_between_lines = 10  # Adjust if needed
    grouped_boxes = []
    current_group = []

    for box in bounding_boxes:
        if not current_group:
            current_group.append(box)
        else:
            min_y = min(current_group, key=lambda x: x[1])[1]
            if box[1] - min_y <= vertical_distance_between_lines:
                current_group.append(box)
            else:
                grouped_boxes.append(current_group)
                current_group = [box]

    if current_group:
        grouped_boxes.append(current_group)

    # Sort left-to-right within each group
    for group in grouped_boxes:
        group.sort(key=lambda box: box[0])

    return grouped_boxes
def sort_bounding_boxes(input_dir, output_dir=None):
    """
    Sorts bounding boxes in all .txt files within a single folder.
    Args:
        input_dir (str or Path): Folder with raw bounding box .txt files.
        output_dir (str or Path, optional): Where to save sorted .txt files. 
            Defaults to '<input_dir>_sorted'.
    """
    input_dir = Path(input_dir)
    if output_dir is None:
        output_dir = input_dir.parent / f"{input_dir.name}_sorted"
    else:
        output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"Processing bounding boxes in: {input_dir}")

    for txt_file in input_dir.glob("*.txt"):
        # Read and clean lines
        with open(txt_file, "r", encoding="utf-8") as f:
            lines = [line for line in f if line.strip()]

        # Write cleaned data to a temp file
        tmp_file = output_dir / f"{txt_file.stem}_tmp.txt"
        tmp_file.write_text("".join(lines))
        sorted_groups = process_bounding_boxes(tmp_file)
        tmp_file.unlink()  # Remove temp file

        # Write sorted result
        output_file = output_dir / f"{txt_file.stem}_sorted.txt"
        with open(output_file, "w", encoding="utf-8") as f_out:
            for group in sorted_groups:
                for box in group:
                    flat = [str(coord) for coord in box]
                    f_out.write(",".join(flat) + "\n")
                f_out.write(";\n")

        print(f"Sorted boxes written: {output_file.name}")

    print(f"All bounding boxes sorted. Output in: {output_dir}")  
def extract_bounding_boxes(image_path, bounding_boxes_file, output_folder, word):
    # Read the main image
    main_image = cv2.imread(image_path)
    # Create the output folder if it doesn't exist
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    if not os.path.exists(image_path):
        return word
    if not os.path.exists(bounding_boxes_file):
        return word

    img = cv2.imread(image_path)
    if img is None:
        return word

    with open(bounding_boxes_file, "r", encoding='utf-8') as f:
        lines = [line.strip() for line in f if line.strip() and line.strip() != ';']
    
    if not lines:
        return word

    # Read bounding box coordinates from the text file
    with open(bounding_boxes_file, 'r', encoding="utf-8") as f:
        bounding_boxes_data = f.read().split(';')
    bounding_boxes_data = bounding_boxes_data[1:]

    img = cv2.imread(image_path)
    if img is None:
        print("Failed to load image.")
        return word

    for indx in range(len(bounding_boxes_data)-1):
        bounding_box_coords = bounding_boxes_data[indx].strip().split('\n')
        for cnt in range(len(bounding_box_coords)):
            coordinates_list = [int(coord) for coord in bounding_box_coords[cnt].split(',')]
            x_min, y_min, x_max, y_min, x_max, y_max, x_min, y_max = coordinates_list

            # Extract the bounding box from the main image
            bounding_box = main_image[y_min:y_max, x_min:x_max]

            # Save the bounding box as a separate image
            output_path = os.path.join(output_folder, f'{word}.png')
            cv2.imwrite(output_path, bounding_box)
            
            word += 1

    return word
def extract_bounding_boxes_train(image_path, bounding_boxes_file, text_file, output_folder, skip_header_groups=1):
    # Load image
    main_image = cv2.imread(image_path)
    os.makedirs(output_folder, exist_ok=True)

    # Load and clean bounding boxes
    with open(bounding_boxes_file, 'r', encoding='utf-8') as f:
        bounding_boxes_data = f.read().split(';')
    bounding_boxes_data = [b.strip() for b in bounding_boxes_data if b.strip()][skip_header_groups:]

    # Load and clean text lines
    with open(text_file, 'r', encoding='utf-8') as f:
        text_data = [line.strip() for line in f if line.strip()]

    # Crop and save word images
    for indx, line in enumerate(text_data):
        if indx >= len(bounding_boxes_data): break

        words = line.split()
        bbox_lines = bounding_boxes_data[indx].splitlines()

        for cnt in range(min(len(words), len(bbox_lines))):
            try:
                coords = list(map(int, bbox_lines[cnt].strip().split(',')))
                x_coords = coords[::2]
                y_coords = coords[1::2]

                x_min, x_max = max(0, min(x_coords)), max(x_coords)
                y_min, y_max = max(0, min(y_coords)), max(y_coords)

                cropped = main_image[y_min:y_max, x_min:x_max]
                if cropped.size == 0:
                    print(f"Empty crop for word '{words[cnt]}' in {image_path}")
                    continue

                output_path = os.path.join(output_folder, f'{words[cnt]}.png')
                cv2.imwrite(output_path, cropped)
            except Exception as e:
                print(f"Failed to process box {cnt} on line {indx} in {bounding_boxes_file}: {e}")

# ------------------ Text Splitting ------------------
def process_textfiles(textfile, sorted_BoundBox_folder, output_folder):
    """
    Splits the ground truth text into per-page files based on 'PDF pX' markers.
    Discards anything before the first 'PDF pX' as disclaimer.
    """
    sorted_BoundBox_folder = Path(sorted_BoundBox_folder)
    output_folder = Path(output_folder)
    output_folder.mkdir(parents=True, exist_ok=True)

    with open(textfile, 'r', encoding='utf-8') as f:
        lines = [line.strip() for line in f if line.strip()]

    # Start collecting pages after the first 'PDF pX'
    pages = []
    current_page = []
    collecting = False

    for line in lines:
        if re.match(r"(?i)^pdf p\d+", line):  # Matches 'PDF pX', case-insensitive
            if collecting:
                pages.append(current_page)
                current_page = []
            else:
                collecting = True  # Start collecting after the first match
        elif collecting:
            current_page.append(line)

    if current_page:
        pages.append(current_page)

    # Sort bounding box files by page number
    bbox_files = sorted(
        sorted_BoundBox_folder.glob('res_image_*_*_sorted.txt'),
        key=lambda f: int(f.name.split('_')[2])
    )

    print(f"Found {len(pages)} text pages and {len(bbox_files)} bbox files.")

    for i, bbox_file in enumerate(bbox_files):
        if i >= len(pages):
            print(f"Not enough text pages for {bbox_file.name}")
            break

        output_file = output_folder / bbox_file.name.replace("_sorted.txt", "_actual.txt")
        with open(output_file, 'w', encoding='utf-8') as output:
            output.write("\n".join(pages[i]) + "\n")

    print(f"Text splitting complete. Output written to: {output_folder}")
def split_text(
    text_path="content/combined_all_text.txt",
    sorted_boxes_folder="content/preprocessing/BoundBoxApplied_sorted",
    output_folder="content/textSplitted"
):
  
    from pathlib import Path
    output_folder = Path(output_folder)
    output_folder.mkdir(parents=True, exist_ok=True)

    print(f"Splitting text for combined dataset:")
    process_textfiles(
        str(text_path),
        str(sorted_boxes_folder),
        str(output_folder)
    )
    print(f"Text splitting complete for combined dataset → {output_folder}")

# ------------------ CSV and Extractions ------------------
def apply_extraction_to_folder_for_train(
    image_folder, 
    bounding_box_folder, 
    text_folder, 
    output_folder,
    exclude_last_n=6  # number of images to skip at end (for test)
):
    """
    Applies extraction to each image, saving bounding boxes into a subfolder per page.
    """
    image_folder = Path(image_folder)
    bounding_box_folder = Path(bounding_box_folder)
    text_folder = Path(text_folder)
    output_folder = Path(output_folder)
    output_folder.mkdir(parents=True, exist_ok=True)

    # Get sorted image list
    image_files = natsorted(image_folder.glob("*.png"))

    if exclude_last_n:
        image_files = image_files[:-exclude_last_n] if len(image_files) > exclude_last_n else []

    for img_path in image_files:
        base = img_path.stem  # e.g. image_1_Nobleza_Virtuosa
        bbox_file = bounding_box_folder / f"res_{base}_sorted.txt"
        text_file = text_folder / f"res_{base}_actual.txt"

        if bbox_file.exists() and text_file.exists():
            # 🔹 Create a dedicated subfolder for this page
            page_output_folder = output_folder / base
            page_output_folder.mkdir(parents=True, exist_ok=True)

            extract_bounding_boxes_train(
                str(img_path),
                str(bbox_file),
                str(text_file),
                str(page_output_folder)
            )
        else:
            if not bbox_file.exists():
                print(f"Bounding box file missing: {bbox_file.name}")
            if not text_file.exists():
                print(f"Text file missing: {text_file.name}")

    print(f"\nAll pages processed. Extracted data saved to: {output_folder}")
def apply_extraction_to_folder_for_test(
    image_folder,
    bounding_box_folder,
    output_folder,
    num_test=6  # number of images to process as test
):
    """
    Applies extraction to the last N images in a combined folder for test purposes.
    Only bounding boxes are required (no text file).
    """
    image_folder = Path(image_folder)
    bounding_box_folder = Path(bounding_box_folder)
    output_folder = Path(output_folder)
    output_folder.mkdir(parents=True, exist_ok=True)

    # List and sort all image files
    image_files = natsorted([f for f in image_folder.glob("*.png")])

    # Select the last N images for test
    if num_test > 0:
        image_files = image_files[-num_test:]

    for img_path in image_files:
        base = img_path.stem
        bbox_file = bounding_box_folder / f"res_{base}_sorted.txt"

        if bbox_file.exists():
            # Test extraction usually doesn't require a text file
            extract_bounding_boxes(
                str(img_path),
                str(bbox_file),
                str(output_folder),
                word=0
            )
        else:
            print(f"Bounding box file missing for {img_path.name}, skipping.")

    print(f"Test extraction complete. Output in: {output_folder}")
def create_csv_from_folder(folder_path, csv_file_path):
    folder_path = Path(folder_path)
    rows = []

    for img_file in folder_path.rglob("*.png"):  # Recursively search subfolders
        label = remove_punctuation(img_file.stem)
        rows.append({
            "FILENAME": str(img_file),   # Full relative path to preserve folder structure
            "IDENTITY": label
        })

    df = pd.DataFrame(rows)
    df.to_csv(csv_file_path, index=False)
    print(f"CSV created at {csv_file_path} with {len(df)} entries.")


